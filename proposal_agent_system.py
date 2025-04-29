import os
import logging
import json
import re
import pandas as pd
import numpy as np
from typing import List, Dict, Any, Optional, Union, Tuple
from pathlib import Path
from datetime import datetime
from dotenv import load_dotenv
from smolagents import CodeAgent, OpenAIServerModel, tool, Tool
from smolagents.utils import truncate_content
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Set up logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger('proposal_agent_system')

# Load environment variables
load_dotenv(override=True)

# Check if API keys are properly set
def check_api_keys():
    """Verify that required API keys are properly set"""
    openai_key = os.getenv("OPENAI_API_KEY")
    
    if not openai_key or openai_key.startswith(('your_', 'sk-your')):
        raise ValueError(
            "OpenAI API key not properly configured. Please set a valid OPENAI_API_KEY in your .env file. "
            "You can get an API key from https://platform.openai.com/account/api-keys"
        )

# ====== PDF PROCESSING TOOLS ======

@tool
def extract_text_from_pdf(pdf_path: str, page_start: int = 0, page_end: int = None) -> str:
    """
    Extract text content from PDF files.
    
    Args:
        pdf_path: Path to the PDF file
        page_start: Starting page (0-indexed)
        page_end: Ending page (None for all pages)
        
    Returns:
        Extracted text content with structure preserved
    """
    try:
        import pypdf
        import io
        
        with open(pdf_path, 'rb') as file:
            pdf_reader = pypdf.PdfReader(file)
            
            # Validate page range
            total_pages = len(pdf_reader.pages)
            if page_start < 0:
                page_start = 0
            if page_end is None or page_end > total_pages:
                page_end = total_pages
                
            # Extract text from pages
            text_content = []
            for i in range(page_start, page_end):
                page = pdf_reader.pages[i]
                text = page.extract_text()
                if text:
                    text_content.append(f"--- Page {i+1} ---\n{text}")
            
            return "\n\n".join(text_content)
            
    except ImportError:
        return "Error: pypdf package not installed. Please install using pip install pypdf"
    except FileNotFoundError:
        return f"Error: File not found at path: {pdf_path}"
    except Exception as e:
        return f"Error extracting text from PDF: {str(e)}"


@tool
def extract_tables_from_pdf(pdf_path: str, page_numbers: List[int] = None) -> str:
    """
    Extract tables from PDF files.
    
    Args:
        pdf_path: Path to the PDF file
        page_numbers: List of page numbers to extract tables from (None for all pages)
        
    Returns:
        JSON string containing extracted tables
    """
    try:
        import tabula
        import pandas as pd
        import json
        
        # Extract tables using tabula
        tables = tabula.read_pdf(
            pdf_path, 
            pages=page_numbers,
            multiple_tables=True
        )
        
        # Convert tables to JSON format
        result = []
        for i, table in enumerate(tables):
            # Convert DataFrame to dictionary
            table_dict = table.to_dict(orient='records')
            result.append({
                "table_index": i,
                "data": table_dict
            })
            
        return json.dumps(result)
            
    except ImportError:
        return "Error: tabula-py package not installed. Please install using pip install tabula-py"
    except FileNotFoundError:
        return f"Error: File not found at path: {pdf_path}"
    except Exception as e:
        return f"Error extracting tables from PDF: {str(e)}"


# ====== DOCUMENT PROCESSING TOOLS ======

@tool
def extract_requirements(text: str) -> str:
    """
    Extract requirements from solicitation text.
    
    Args:
        text: Text content from the solicitation
        
    Returns:
        JSON string with extracted requirements
    """
    try:
        import re
        import json
        
        # Patterns for requirements
        requirement_patterns = [
            r"(?:^|\n)(?:.*?)(shall|must|will be required to|is required to|are required to|will|is to)(?:.*?)(?:$|\n)",
            r"(?:^|\n)(?:.*?)(required|requirement|mandatory|essential)(?:.*?)(?:$|\n)"
        ]
        
        # Extract requirements using patterns
        requirements = []
        
        for pattern in requirement_patterns:
            matches = re.finditer(pattern, text, re.IGNORECASE | re.MULTILINE)
            for match in matches:
                req_text = match.group(0).strip()
                requirements.append({
                    "text": req_text,
                    "type": "mandatory" if any(word in req_text.lower() for word in ["shall", "must", "required", "mandatory"]) else "desirable"
                })
        
        # Remove duplicates while preserving order
        unique_requirements = []
        seen = set()
        for req in requirements:
            if req["text"] not in seen:
                seen.add(req["text"])
                unique_requirements.append(req)
        
        return json.dumps({
            "count": len(unique_requirements),
            "requirements": unique_requirements
        })
        
    except Exception as e:
        return json.dumps({
            "error": f"Error extracting requirements: {str(e)}",
            "requirements": []
        })


@tool
def extract_evaluation_criteria(text: str) -> str:
    """
    Extract evaluation criteria from solicitation text.
    
    Args:
        text: Text content from the solicitation
        
    Returns:
        JSON string with extracted evaluation criteria
    """
    try:
        import re
        import json
        
        # Look for sections related to evaluation
        eval_section_patterns = [
            r"(?:^|\n)[\d\.\s]*(?:evaluation criteria|evaluation factors|basis for award|source selection)(?:.*?)(?:\n\n|\n\d|\Z)",
            r"(?:^|\n)[\d\.\s]*Section [A-Z](?:.*)(?:Evaluation|Award)(?:.*?)(?:\n\n|\n\d|\Z)"
        ]
        
        # Extract evaluation criteria sections
        eval_sections = []
        for pattern in eval_section_patterns:
            matches = re.finditer(pattern, text, re.IGNORECASE | re.MULTILINE | re.DOTALL)
            for match in matches:
                eval_sections.append(match.group(0).strip())
        
        # If no sections found, return empty
        if not eval_sections:
            return json.dumps({
                "count": 0,
                "criteria": []
            })
        
        # Now extract the individual criteria from these sections
        criteria = []
        
        # Patterns for finding individual criteria
        criteria_patterns = [
            r"(?:^|\n)[\d\.\s]*(?:Factor|Criterion|Area|Element)\s*(?:#|No|Number)?\s*\d+[:\.\s]+([^\n]+)",
            r"(?:^|\n)[\d\.\s]*([A-Za-z][\w\s\-]+?(?:Technical|Management|Experience|Past Performance|Price|Cost)[^\n]+)"
        ]
        
        # Find criteria in the eval sections
        for section in eval_sections:
            for pattern in criteria_patterns:
                matches = re.finditer(pattern, section, re.IGNORECASE | re.MULTILINE)
                for match in matches:
                    criteria_name = match.group(1).strip() if len(match.groups()) > 0 else match.group(0).strip()
                    
                    # Try to determine weight if mentioned
                    weight_match = re.search(r"(?:weight|importance|value)[^\d]*?(\d+)%", match.group(0), re.IGNORECASE)
                    weight = weight_match.group(1) if weight_match else "Not specified"
                    
                    criteria.append({
                        "name": criteria_name,
                        "weight": weight,
                        "section_text": match.group(0).strip()
                    })
        
        # Remove duplicates
        unique_criteria = []
        seen_names = set()
        for crit in criteria:
            if crit["name"] not in seen_names:
                seen_names.add(crit["name"])
                unique_criteria.append(crit)
        
        return json.dumps({
            "count": len(unique_criteria),
            "criteria": unique_criteria
        })
        
    except Exception as e:
        return json.dumps({
            "error": f"Error extracting evaluation criteria: {str(e)}",
            "criteria": []
        })


@tool
def identify_document_sections(text: str) -> str:
    """
    Identify and extract section structure from document text.
    
    Args:
        text: Text content from the document
        
    Returns:
        JSON string with document sections
    """
    try:
        import re
        import json
        
        # Regular expressions to identify section headers
        section_patterns = [
            # Federal solicitation standard section headers pattern
            r"(?:^|\n)(?:SECTION|PART)\s+([A-Z])\s*[–-]\s*(.+?)(?=\n)",
            # Section with numbers pattern (1.0, 1.1, etc.)
            r"(?:^|\n)(\d+(?:\.\d+)*)\s+(.+?)(?=\n)",
            # Uppercase headers
            r"(?:^|\n)([A-Z][A-Z\s\d]+[A-Z\d])(?:\n|:)",
        ]
        
        sections = []
        
        # Find all sections using patterns
        for pattern in section_patterns:
            matches = re.finditer(pattern, text, re.MULTILINE)
            for match in matches:
                if len(match.groups()) == 2:  # Number/letter + title
                    section_id = match.group(1).strip()
                    section_title = match.group(2).strip()
                else:  # Just a title
                    section_id = ""
                    section_title = match.group(1).strip()
                
                # Calculate the approximate location in the document
                location = match.start() / len(text) if len(text) > 0 else 0
                
                sections.append({
                    "id": section_id,
                    "title": section_title,
                    "full_header": match.group(0).strip(),
                    "location": f"{location:.2%} through document",
                })
        
        # Sort sections by their appearance in the document
        sections.sort(key=lambda x: text.find(x["full_header"]))
        
        return json.dumps({
            "count": len(sections),
            "sections": sections
        })
        
    except Exception as e:
        return json.dumps({
            "error": f"Error identifying document sections: {str(e)}",
            "sections": []
        })


class WebpageVisitorTool(Tool):
    name = "visit_webpage"
    description = (
        "Visits a webpage at the given URL and reads its content as a markdown string. "
        "Use this to browse agency websites, solicitation listings, or reference materials."
    )
    inputs = {
        "url": {
            "type": "string",
            "description": "The URL of the webpage to visit.",
        }
    }
    output_type = "string"

    def forward(self, url: str) -> str:
        try:
            import re
            import requests
            from markdownify import markdownify
            from requests.exceptions import RequestException
            from smolagents.utils import truncate_content
        except ImportError as e:
            raise ImportError(
                "You must install packages `markdownify` and `requests` to run this tool: "
                "run `pip install markdownify requests`."
            ) from e
        
        try:
            # Add headers to mimic a browser request (to avoid being blocked)
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36',
                'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8',
                'Accept-Language': 'en-US,en;q=0.5',
                'Connection': 'keep-alive',
                'Upgrade-Insecure-Requests': '1',
                'Cache-Control': 'max-age=0'
            }
            
            response = requests.get(url, headers=headers, timeout=20)
            response.raise_for_status()
            
            content = response.text
            
            # Check if it's a PDF and handle accordingly
            if 'application/pdf' in response.headers.get('Content-Type', ''):
                return "This is a PDF document. Please use the PDF extraction tools to analyze this content."
            
            markdown_content = markdownify(content).strip()
            markdown_content = re.sub(r"\n{3,}", "\n\n", markdown_content)
            
            # Prioritize content based on proposal relevance
            proposal_markers = ["solicitation", "request for proposal", "RFP", "RFQ", "contract", 
                             "federal", "requirement", "submission", "evaluation", "criteria"]
            
            for marker in proposal_markers:
                pattern = re.compile(f'({marker})', re.IGNORECASE)
                markdown_content = pattern.sub(r'**\1**', markdown_content)
            
            return truncate_content(markdown_content, 40000)

        except requests.exceptions.Timeout:
            return "The request to the resource timed out. Please try again later or check the URL."
        except RequestException as e:
            return f"Error fetching the webpage: {str(e)}"
        except Exception as e:
            return f"An unexpected error occurred when accessing the resource: {str(e)}"


# ====== DOCUMENT GENERATION TOOLS ======

@tool
def generate_docx_outline(title: str, sections: str, output_path: str) -> str:
    """
    Generate a professional DOCX proposal outline with proper formatting.
    
    Args:
        title: Proposal title
        sections: JSON string containing section information
        output_path: Path to save the generated DOCX file
        
    Returns:
        Status message
    """
    try:
        import json
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        # Parse the sections JSON
        sections_data = json.loads(sections)
        
        # Create a new Document
        doc = Document()
        
        # Configure page setup
        sections = doc.sections
        for section in sections:
            section.page_height = Inches(11)
            section.page_width = Inches(8.5)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.header_distance = Inches(0.5)
            section.footer_distance = Inches(0.5)
        
        # Define styles
        styles = doc.styles
        
        # Title style
        title_style = styles.add_style('ProposalTitle', WD_STYLE_TYPE.PARAGRAPH)
        title_style.font.name = 'Arial'
        title_style.font.size = Pt(16)
        title_style.font.bold = True
        title_style.font.color.rgb = RGBColor(0, 51, 102)  # Dark blue
        title_paragraph_format = title_style.paragraph_format
        title_paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_paragraph_format.space_before = Pt(0)
        title_paragraph_format.space_after = Pt(12)
        title_paragraph_format.keep_together = True
        
        # Heading 1 style
        h1_style = styles.add_style('ProposalHeading1', WD_STYLE_TYPE.PARAGRAPH)
        h1_style.font.name = 'Arial'
        h1_style.font.size = Pt(14)
        h1_style.font.bold = True
        h1_style.font.color.rgb = RGBColor(0, 51, 102)  # Dark blue
        h1_paragraph_format = h1_style.paragraph_format
        h1_paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        h1_paragraph_format.space_before = Pt(12)
        h1_paragraph_format.space_after = Pt(6)
        h1_paragraph_format.keep_with_next = True
        
        # Heading 2 style
        h2_style = styles.add_style('ProposalHeading2', WD_STYLE_TYPE.PARAGRAPH)
        h2_style.font.name = 'Arial'
        h2_style.font.size = Pt(12)
        h2_style.font.bold = True
        h2_style.font.color.rgb = RGBColor(0, 51, 102)  # Dark blue
        h2_paragraph_format = h2_style.paragraph_format
        h2_paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        h2_paragraph_format.space_before = Pt(10)
        h2_paragraph_format.space_after = Pt(6)
        h2_paragraph_format.keep_with_next = True
        
        # Body text style
        body_style = styles.add_style('ProposalBody', WD_STYLE_TYPE.PARAGRAPH)
        body_style.font.name = 'Arial'
        body_style.font.size = Pt(11)
        body_paragraph_format = body_style.paragraph_format
        body_paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        body_paragraph_format.space_before = Pt(6)
        body_paragraph_format.space_after = Pt(6)
        
        # Add title
        doc.add_paragraph(title, style='ProposalTitle')
        
        # Add sections
        current_level = 0
        section_counter = [0, 0, 0]  # Track section numbering up to 3 levels
        
        for section in sections_data:
            # Get section level (0, 1, or 2) based on format or ID structure
            level = 0
            if 'level' in section:
                level = min(section['level'], 2)  # Max level 2 (0, 1, 2)
            elif 'id' in section and '.' in section['id']:
                # Count the dots to determine level (e.g., "1.2.3" = level 2)
                level = min(section['id'].count('.'), 2)
            
            # Update section counters
            if level == 0:
                section_counter[0] += 1
                section_counter[1] = 0
                section_counter[2] = 0
            elif level == 1:
                section_counter[1] += 1
                section_counter[2] = 0
            else:  # level == 2
                section_counter[2] += 1
            
            # Format section number
            if level == 0:
                section_number = f"{section_counter[0]}."
            elif level == 1:
                section_number = f"{section_counter[0]}.{section_counter[1]}"
            else:  # level == 2
                section_number = f"{section_counter[0]}.{section_counter[1]}.{section_counter[2]}"
            
            # Add section with appropriate style
            if level == 0:
                p = doc.add_paragraph(style='ProposalHeading1')
                p.add_run(f"{section_number} {section['title']}")
            elif level == 1:
                p = doc.add_paragraph(style='ProposalHeading2')
                p.add_run(f"{section_number} {section['title']}")
            else:  # level == 2
                p = doc.add_paragraph(style='ProposalBody')
                p.add_run(f"{section_number} {section['title']}").bold = True
            
            # Add placeholder text if it exists
            if 'placeholder' in section and section['placeholder']:
                body_text = doc.add_paragraph(style='ProposalBody')
                body_text.add_run(section['placeholder'])
        
        # Save the document
        doc.save(output_path)
        
        return f"Successfully generated proposal outline DOCX at {output_path}"
        
    except ImportError:
        return "Error: python-docx package not installed. Please install using pip install python-docx"
    except json.JSONDecodeError:
        return "Error: Invalid JSON structure for sections"
    except Exception as e:
        return f"Error generating DOCX outline: {str(e)}"


@tool
def add_compliance_matrix(docx_path: str, requirements: str, output_path: str = None) -> str:
    """
    Add a compliance matrix to an existing DOCX document.
    
    Args:
        docx_path: Path to the existing DOCX document
        requirements: JSON string containing requirements
        output_path: Path to save the updated DOCX file (if None, overwrites original)
        
    Returns:
        Status message
    """
    try:
        import json
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        # Parse the requirements JSON
        requirements_data = json.loads(requirements)
        
        # Load the document
        doc = Document(docx_path)
        
        # Add a section break
        doc.add_page_break()
        
        # Add compliance matrix title
        title = doc.add_paragraph("Compliance Matrix", style='ProposalTitle')
        
        # Add introduction text
        intro = doc.add_paragraph(style='ProposalBody')
        intro.add_run("This compliance matrix maps each solicitation requirement to the relevant section(s) of our proposal response.").italic = True
        
        # Create table for compliance matrix
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        
        # Set column widths
        table.autofit = False
        table.columns[0].width = Inches(0.5)  # Item #
        table.columns[1].width = Inches(3.5)  # Requirement
        table.columns[2].width = Inches(1.5)  # Section Reference
        table.columns[3].width = Inches(1.0)  # Compliant (Y/N)
        
        # Add header row
        header_cells = table.rows[0].cells
        header_cells[0].text = "Item #"
        header_cells[1].text = "Requirement"
        header_cells[2].text = "Proposal Section"
        header_cells[3].text = "Compliant"
        
        # Style the header row
        for cell in header_cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in cell.paragraphs[0].runs:
                run.font.bold = True
                run.font.size = Pt(11)
        
        # Add requirements to the table
        for i, req in enumerate(requirements_data.get('requirements', [])):
            row = table.add_row()
            cells = row.cells
            
            # Item number
            cells[0].text = str(i+1)
            cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Requirement text
            cells[1].text = req['text']
            
            # Placeholder for section reference (to be filled in)
            cells[2].text = "TBD"
            cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Compliance indicator
            cells[3].text = "☐"
            cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Save the document
        if output_path:
            doc.save(output_path)
        else:
            doc.save(docx_path)  # Overwrite original
        
        save_path = output_path if output_path else docx_path
        return f"Successfully added compliance matrix to DOCX at {save_path}"
        
    except ImportError:
        return "Error: python-docx package not installed. Please install using pip install python-docx"
    except json.JSONDecodeError:
        return "Error: Invalid JSON structure for requirements"
    except Exception as e:
        return f"Error adding compliance matrix: {str(e)}"


# ====== PROPOSAL AGENT SYSTEM ======

# 1. Document Processing Agent
def create_document_processing_agent(model_name="gpt-4.1-mini"):
    """Create a document processing agent specialized for federal solicitations"""
    processing_agent = CodeAgent(
        model=OpenAIServerModel(
            model_name,
            max_completion_tokens=8096,
        ),
        tools=[extract_text_from_pdf, extract_tables_from_pdf, identify_document_sections, WebpageVisitorTool()],
        max_steps=15,
        name="document_processing_agent",
        description="Processes federal solicitation documents, extracts structured content, and organizes information."
    )
    processing_agent.logger.console.width = 100
    return processing_agent

# 2. Requirements Analysis Agent
def create_requirements_analysis_agent(model_name="gpt-4.1-mini"):
    """Create a requirements analysis agent specialized for federal solicitations"""
    requirements_agent = CodeAgent(
        model=OpenAIServerModel(
            model_name,
            max_completion_tokens=8096,
        ),
        tools=[extract_requirements, extract_evaluation_criteria, WebpageVisitorTool()],
        max_steps=10,
        name="requirements_analysis_agent",
        description="Analyzes solicitation requirements, identifies evaluation criteria, and creates compliance matrices."
    )
    requirements_agent.logger.console.width = 100
    return requirements_agent

# 3. Proposal Development Agent
def create_proposal_development_agent(model_name="gpt-4.1-mini"):
    """Create a proposal development agent for creating formatted proposal documents"""
    proposal_agent = CodeAgent(
        model=OpenAIServerModel(
            model_name,
            max_completion_tokens=8096,
        ),
        tools=[generate_docx_outline, add_compliance_matrix],
        max_steps=10,
        name="proposal_development_agent",
        description="Creates professionally formatted proposal outlines, response documents, and compliance materials."
    )
    proposal_agent.logger.console.width = 100
    return proposal_agent

# 4. Validation function for proposal document
def validate_proposal_document(final_answer, agent_memory):
    """
    Validates that the generated proposal document meets quality standards:
    - Properly formatted DOCX with professional styling
    - Complete structure matching solicitation requirements
    - Includes all required sections
    - Has compliance matrix if needed
    """
    docx_path = "proposal_outline.docx"
    
    # Check if the document exists
    if not os.path.exists(docx_path):
        raise Exception("Proposal document was not created. Expected 'proposal_outline.docx'")
    
    try:
        # Load the document to validate
        from docx import Document
        doc = Document(docx_path)
        
        # Basic validation
        if len(doc.paragraphs) < 5:
            raise Exception("Document is too short - missing content")
        
        # Check for basic structure (you could make this more comprehensive)
        has_title = False
        has_headings = False
        has_sections = False
        
        for para in doc.paragraphs:
            if para.style.name == 'ProposalTitle':
                has_title = True
            if para.style.name == 'ProposalHeading1' or para.style.name == 'Heading 1':
                has_headings = True
                has_sections = True
        
        if not has_title:
            raise Exception("Document is missing a title")
        
        if not has_headings:
            raise Exception("Document is missing section headings")
        
        if not has_sections:
            raise Exception("Document has improper structure - missing proper sections")
    
        return True
    
    except Exception as e:
        raise Exception(f"Proposal document validation failed: {str(e)}")

# 5. Manager Agent for Proposal Development
def create_proposal_manager_agent(processing_agent, requirements_agent, proposal_agent, model_name="gpt-4.1-mini"):
    """Create a manager agent that coordinates the proposal development process"""
    
    # Create the manager agent with appropriate tools and sub-agents
    manager_agent = CodeAgent(
        model=OpenAIServerModel(
            model_name,
            max_tokens=8096,
        ),
        tools=[],
        managed_agents=[processing_agent, requirements_agent, proposal_agent],
        additional_authorized_imports=[
            "docx",
            "docx.shared",
            "docx.enum.text",
            "docx.enum.style",
            "docx.oxml",
            "os",
            "json",
            "re",
            "pandas",
            "pathlib"
        ],
        planning_interval=5,
        verbosity_level=2,
        final_answer_checks=[validate_proposal_document],
        max_steps=30,
    )
    
    manager_agent.logger.console.width = 100
    return manager_agent

# ====== USAGE FUNCTION ======

def run_proposal_agent(solicitation_path, output_path="proposal_outline.docx"):
    """
    Run the proposal agent system to analyze a solicitation and create a proposal outline.
    
    Args:
        solicitation_path: Path to the solicitation PDF file
        output_path: Path to save the generated proposal document
    
    Returns:
        Result message
    """
    
    # Check API keys before proceeding
    try:
        check_api_keys()
    except ValueError as e:
        logger.error(f"API key error: {str(e)}")
        return str(e)
    
    # Set up the model name - you can change this based on available models
    model_name = os.getenv("DEFAULT_MODEL", "gpt-4.1-mini")
    
    try:
        # Ensure solicitation file exists
        if not os.path.exists(solicitation_path):
            return f"Error: Solicitation file not found at {solicitation_path}"
        
        # Create the agent hierarchy
        logger.info("Setting up Proposal Agent System...")
        processing_agent = create_document_processing_agent(model_name)
        requirements_agent = create_requirements_analysis_agent(model_name)
        proposal_agent = create_proposal_development_agent(model_name)
        
        # Create the manager agent
        manager_agent = create_proposal_manager_agent(
            processing_agent, requirements_agent, proposal_agent, model_name
        )
        
        # Remove previous output file if it exists
        if os.path.exists(output_path):
            os.remove(output_path)
        
        # Prepare the query for the agent
        query = f"""
        Analyze the federal solicitation document at "{solicitation_path}" and create a professional proposal outline document.
        
        Follow these steps:
        1. Process and extract all text from the solicitation PDF
        2. Identify key sections and structure of the solicitation
        3. Extract all requirements and evaluation criteria
        4. Create a proposal outline that follows the solicitation structure
        5. Include all required sections and appropriate headings
        6. Add a compliance matrix that maps to all requirements
        7. Format the document professionally using proper styles
        8. Save the final proposal outline as "{output_path}"
        
        Use advanced DOCX formatting features including:
        - Professional proposal title and heading styles
        - Consistent paragraph formatting and spacing
        - Proper section numbering and hierarchy
        - Compliance matrix with requirements traceability
        
        Make sure the final document is well-structured, professionally formatted, and addresses all requirements in the solicitation.
        """
        
        # Run the proposal agent
        logger.info(f"Starting Proposal Development for: {solicitation_path}")
        result = manager_agent.run(query)
        
        logger.info("\n===== Proposal Development Completed =====")
        return result
    
    except Exception as e:
        import traceback
        logger.error(f"\n===== ERROR DURING PROPOSAL DEVELOPMENT =====")
        logger.error(f"Error type: {type(e).__name__}")
        logger.error(f"Error message: {str(e)}")
        logger.error("\nStack trace:")
        traceback.print_exc()
        
        return f"Proposal development failed: {str(e)}"

# Example usage
if __name__ == "__main__":
    # Path to your federal solicitation PDF
    solicitation_path = "sample_solicitation.pdf"
    
    # Run the proposal agent
    result = run_proposal_agent(solicitation_path)
    print(result)
